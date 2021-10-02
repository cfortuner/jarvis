from bs4 import BeautifulSoup
import re
from html.entities import name2codepoint

from builtins import str, bytes, dict, int
from builtins import map, zip, filter
from builtins import object

from html.parser import HTMLParser as _HTMLParser
import unicodedata

# Python 3
from urllib.parse import (
    urlparse,
    urljoin,
    urlsplit,
    urlencode,
    quote_plus,
    unquote_plus,
)
from urllib.request import (
    urlopen,
    Request,
    HTTPHandler,
    ProxyHandler,
    HTTPCookieProcessor,
    install_opener,
    build_opener,
)
from urllib.error import HTTPError as UrllibHTTPError
from urllib.error import URLError as UrllibURLError

# --- STRING FUNCTIONS ------------------------------------------------------------------------------
# Latin-1 (ISO-8859-1) encoding is identical to Windows-1252 except for the code points 128-159:
# Latin-1 assigns control codes in this range, Windows-1252 has characters, punctuation, symbols
# assigned to these code points.


def decode_string(v, encoding="utf-8"):
    """Returns the given value as a Unicode string (if possible)."""
    if isinstance(encoding, str):
        encoding = ((encoding,),) + (("windows-1252",), ("utf-8", "ignore"))
    if isinstance(v, bytes):
        for e in encoding:
            try:
                return v.decode(*e)
            except:
                pass
        return v
    return str(v)


def encode_string(v, encoding="utf-8"):
    """Returns the given value as a Python byte string (if possible)."""
    if isinstance(encoding, str):
        encoding = ((encoding,),) + (("windows-1252",), ("utf-8", "ignore"))
    if isinstance(v, str):
        for e in encoding:
            try:
                return v.encode(*e)
            except:
                pass
        return v
    return bytes(v)


#### PATTERN | WEB #################################################################################
# -*- coding: utf-8 -*-
# Copyright (c) 2010 University of Antwerp, Belgium
# Author: Tom De Smedt <tom@organisms.be>
# License: BSD (see LICENSE.txt for details).
# http://www.clips.ua.ac.be/pages/pattern

####################################################################################################
# Python API interface for various web services (Google, Twitter, Wikipedia, ...)


from builtins import str, bytes, dict, int, chr
from builtins import map, filter, zip
from builtins import object

import os
import sys
import threading
import time
import socket

from io import open

# Python 3
from urllib.parse import (
    urlparse,
    urljoin,
    urlsplit,
    urlencode,
    quote_plus,
    unquote_plus,
)
from urllib.request import (
    HTTPCookieProcessor,
    HTTPHandler,
    ProxyHandler,
    Request,
    build_opener,
    install_opener,
    urlopen,
)
from urllib.error import HTTPError as UrllibHTTPError
from urllib.error import URLError as UrllibURLError

import base64

from html.entities import name2codepoint

# Python 3
import http.client as httplib

from html.parser import HTMLParser as _HTMLParser

import re
import unicodedata
from io import StringIO
import bisect
import itertools

# Python 3: We don't actually need it (in this case)
new = None


import bs4 as BeautifulSoup

cache = {}

try:
    MODULE = os.path.dirname(os.path.realpath(__file__))
except:
    MODULE = ""

#### UNICODE #######################################################################################
# Latin-1 (ISO-8859-1) encoding is identical to Windows-1252 except for the code points 128-159:
# Latin-1 assigns control codes in this range, Windows-1252 has characters, punctuation, symbols
# assigned to these code points.


u = decode_utf8 = decode_string
s = encode_utf8 = encode_string

GREMLINS = set(
    [
        0x0152,
        0x0153,
        0x0160,
        0x0161,
        0x0178,
        0x017E,
        0x017D,
        0x0192,
        0x02C6,
        0x02DC,
        0x2013,
        0x2014,
        0x201A,
        0x201C,
        0x201D,
        0x201E,
        0x2018,
        0x2019,
        0x2020,
        0x2021,
        0x2022,
        0x2026,
        0x2030,
        0x2039,
        0x203A,
        0x20AC,
        0x2122,
    ]
)


def fix(s, ignore=""):
    """Returns a Unicode string that fixes common encoding problems (Latin-1, Windows-1252).
    For example: fix("clichÃ©") => "cliché".
    """
    # http://blog.luminoso.com/2012/08/20/fix-unicode-mistakes-with-python/
    if not isinstance(s, str):
        s = s.decode("utf-8")
        # If this doesn't work,
        # copy & paste string in a Unicode .txt,
        # and then pass open(f).read() to fix().
    u = []
    i = 0
    for j, ch in enumerate(s):
        if ch in ignore:
            continue
        if ord(ch) < 128:  # ASCII
            continue
        if ord(ch) in GREMLINS:
            ch = ch.encode("windows-1252")
        else:
            try:
                ch = ch.encode("latin-1")
            except:
                ch = ch.encode("utf-8")
        u.append(s[i:j].encode("utf-8"))
        u.append(ch)
        i = j + 1
    u.append(s[i:].encode("utf-8"))
    u = b"".join(u)
    u = u.decode("utf-8", "replace")
    u = u.replace("\n", "\n ")
    u = u.split(" ")
    # Revert words that have the replacement character,
    # i.e., fix("cliché") should not return "clich�".
    for i, (w1, w2) in enumerate(zip(s.split(" "), u)):
        if "\ufffd" in w2:  # �
            u[i] = w1
    u = " ".join(u)
    u = u.replace("\n ", "\n")
    return u


def latin(s):
    """Returns True if the string contains only Latin-1 characters
    (no Chinese, Japanese, Arabic, Cyrillic, Hebrew, Greek, ...).
    """
    if not isinstance(s, str):
        s = s.decode("utf-8")
    return all(unicodedata.name(ch).startswith("LATIN") for ch in s if ch.isalpha())


# For clearer source code:
bytestring = b = s

#### ASYNCHRONOUS REQUEST ##########################################################################


class AsynchronousRequest(object):
    def __init__(self, function, *args, **kwargs):
        """Executes the function in the background.
        AsynchronousRequest.done is False as long as it is busy, but the program will not halt in the meantime.
        AsynchronousRequest.value contains the function's return value once done.
        AsynchronousRequest.error contains the Exception raised by an erronous function.
        For example, this is useful for running live web requests while keeping an animation running.
        For good reasons, there is no way to interrupt a background process (i.e. Python thread).
        You are responsible for ensuring that the given function doesn't hang.
        """
        self._response = None  # The return value of the given function.
        self._error = None  # The exception (if any) raised by the function.
        self._time = time.time()
        self._function = function
        self._thread = threading.Thread(
            target=self._fetch, args=(function,) + args, kwargs=kwargs
        )
        self._thread.start()

    def _fetch(self, function, *args, **kwargs):
        """Executes the function and sets AsynchronousRequest.response."""
        try:
            self._response = function(*args, **kwargs)
        except Exception as e:
            self._error = e

    def now(self):
        """Waits for the function to finish and yields its return value."""
        self._thread.join()
        return self._response

    @property
    def elapsed(self):
        return time.time() - self._time

    @property
    def done(self):
        return not self._thread.isAlive()

    @property
    def value(self):
        return self._response

    @property
    def error(self):
        return self._error

    def __repr__(self):
        return "AsynchronousRequest(function='%s')" % self._function.__name__


def asynchronous(function, *args, **kwargs):
    """Returns an AsynchronousRequest object for the given function."""
    return AsynchronousRequest(function, *args, **kwargs)


send = asynchronous

#### URL ###########################################################################################

# User agent and referrer.
# Used to identify the application accessing the web.
USER_AGENT = "Pattern/2.6 +http://www.clips.ua.ac.be/pattern"
REFERRER = "http://www.clips.ua.ac.be/pattern"

# Mozilla user agent.
# Websites can include code to block out any application except browsers.
MOZILLA = "Mozilla/5.0"

# HTTP request method.
GET = "get"  # Data is encoded in the URL.
POST = "post"  # Data is encoded in the message body.

# URL parts.
# protocol://username:password@domain:port/path/page?query_string#anchor
PROTOCOL, USERNAME, PASSWORD, DOMAIN, PORT, PATH, PAGE, QUERY, ANCHOR = (
    "protocol",
    "username",
    "password",
    "domain",
    "port",
    "path",
    "page",
    "query",
    "anchor",
)

# MIME type.
MIMETYPE_WEBPAGE = ["text/html"]
MIMETYPE_STYLESHEET = ["text/css"]
MIMETYPE_PLAINTEXT = ["text/plain"]
MIMETYPE_PDF = ["application/pdf"]
MIMETYPE_NEWSFEED = ["application/rss+xml", "application/atom+xml"]
MIMETYPE_IMAGE = ["image/gif", "image/jpeg", "image/png", "image/tiff"]
MIMETYPE_AUDIO = ["audio/mpeg", "audio/mp4", "audio/x-aiff", "audio/x-wav"]
MIMETYPE_VIDEO = [
    "video/mpeg",
    "video/mp4",
    "video/avi",
    "video/quicktime",
    "video/x-flv",
]
MIMETYPE_ARCHIVE = ["application/x-stuffit", "application/x-tar", "application/zip"]
MIMETYPE_SCRIPT = ["application/javascript", "application/ecmascript"]


def extension(filename):
    """Returns the extension in the given filename: "cat.jpg" => ".jpg"."""
    return os.path.splitext(filename)[1]


def urldecode(query):
    """Inverse operation of urlencode.
    Returns a dictionary of (name, value)-items from a URL query string.
    """

    def _format(s):
        if s == "" or s == "None":
            return None
        if s.lstrip("-").isdigit():
            return int(s)
        try:
            return float(s)
        except:
            return s

    if query:
        query = query.lstrip("?").split("&")
        query = ((kv.split("=") + [None])[:2] for kv in query)
        if sys.version > "3":
            query = (
                (u(unquote_plus(k)), _format(u(unquote_plus(v))))
                for k, v in query
                if k != ""
            )
        else:
            query = (
                (
                    u(unquote_plus(bytestring(k))),
                    _format(u(unquote_plus(bytestring(v)))),
                )
                for k, v in query
                if k != ""
            )
        return dict(query)
    return {}


url_decode = urldecode


def proxy(host, protocol="https"):
    """Returns the value for the URL.open() proxy parameter.
    - host: host address of the proxy server.
    """
    return (host, protocol)


class Error(Exception):
    """Base class for pattern.web errors."""

    def __init__(self, *args, **kwargs):
        Exception.__init__(self, *args)
        self.src = kwargs.pop("src", None)
        self.url = kwargs.pop("url", None)

    @property
    def headers(self):
        return dict(list(self.src.headers.items()))


class URLError(Error):
    pass  # URL contains errors (e.g. a missing t in htp://).


class URLTimeout(URLError):
    pass  # URL takes to long to load.


class HTTPError(URLError):
    pass  # URL causes an error on the contacted server.


class HTTP301Redirect(HTTPError):
    pass  # Too many redirects.
    # The site may be trying to set a cookie and waiting for you to return it,
    # or taking other measures to discern a browser from a script.
    # For specific purposes you should build your own urllib2.HTTPRedirectHandler
    # and pass it to urllib2.build_opener() in URL.open()


class HTTP400BadRequest(HTTPError):
    pass  # URL contains an invalid request.


class HTTP401Authentication(HTTPError):
    pass  # URL requires a login and password.


class HTTP403Forbidden(HTTPError):
    pass  # URL is not accessible (user-agent?)


class HTTP404NotFound(HTTPError):
    pass  # URL doesn't exist on the internet.


class HTTP414RequestURITooLong(HTTPError):
    pass  # URL is too long.


class HTTP420Error(HTTPError):
    pass  # Used by Twitter for rate limiting.


class HTTP429TooMayRequests(HTTPError):
    pass  # Used by Twitter for rate limiting.


class HTTP500InternalServerError(HTTPError):
    pass  # Generic server error.


class HTTP503ServiceUnavailable(HTTPError):
    pass  # Used by Bing for rate limiting.


class URL(object):
    def __init__(self, string="", method=GET, query={}, **kwargs):
        """URL object with the individual parts available as attributes:
        For protocol://username:password@domain:port/path/page?query_string#anchor:
        - URL.protocol: http, https, ftp, ...
        - URL.username: username for restricted domains.
        - URL.password: password for restricted domains.
        - URL.domain  : the domain name, e.g. nodebox.net.
        - URL.port    : the server port to connect to.
        - URL.path    : the server path of folders, as a list, e.g. ['news', '2010']
        - URL.page    : the page name, e.g. page.html.
        - URL.query   : the query string as a dictionary of (name, value)-items.
        - URL.anchor  : the page anchor.
        If method is POST, the query string is sent with HTTP POST.
        """
        self.__dict__[
            "method"
        ] = method  # Use __dict__ directly since __setattr__ is overridden.
        self.__dict__["_string"] = u(string)
        self.__dict__["_parts"] = None
        self.__dict__["_headers"] = None
        self.__dict__["_redirect"] = None
        if isinstance(string, URL):
            self.__dict__["method"] = string.method
            self.query.update(string.query)
        if len(query) > 0:
            # Requires that we parse the string first (see URL.__setattr__).
            self.query.update(query)
        if len(kwargs) > 0:
            # Requires that we parse the string first (see URL.__setattr__).
            self.parts.update(kwargs)

    def _parse(self):
        """Parses all the parts of the URL string to a dictionary.
        URL format: protocal://username:password@domain:port/path/page?querystring#anchor
        For example: http://user:pass@example.com:992/animal/bird?species=seagull&q#wings
        This is a cached method that is only invoked when necessary, and only once.
        """
        p = urlsplit(self._string)
        P = {
            PROTOCOL: p[0],  # http
            USERNAME: "",  # user
            PASSWORD: "",  # pass
            DOMAIN: p[1],  # example.com
            PORT: "",  # 992
            PATH: p[2],  # [animal]
            PAGE: "",  # bird
            QUERY: urldecode(p[3]),  # {"species": "seagull", "q": None}
            ANCHOR: p[4],  # wings
        }

        # Split the username and password from the domain.
        if "@" in P[DOMAIN]:
            P[USERNAME], P[PASSWORD] = (p[1].split("@")[0].split(":") + [""])[:2]
            P[DOMAIN] = p[1].split("@")[1]
        # Split the port number from the domain.
        if ":" in P[DOMAIN]:
            P[DOMAIN], P[PORT] = P[DOMAIN].split(":")
            P[PORT] = P[PORT].isdigit() and int(P[PORT]) or P[PORT]
        # Split the base page from the path.
        if "/" in P[PATH]:
            P[PAGE] = p[2].split("/")[-1]
            P[PATH] = p[2][: len(p[2]) - len(P[PAGE])].strip("/").split("/")
            P[PATH] = list(filter(lambda v: v != "", P[PATH]))
        else:
            P[PAGE] = p[2].strip("/")
            P[PATH] = []
        self.__dict__["_parts"] = P

    # URL.string yields unicode(URL) by joining the different parts,
    # if the URL parts have been modified.
    def _get_string(self):
        return str(self)

    def _set_string(self, v):
        self.__dict__["_string"] = u(v)
        self.__dict__["_parts"] = None

    string = property(_get_string, _set_string)

    @property
    def parts(self):
        """Yields a dictionary with the URL parts."""
        if not self._parts:
            self._parse()
        return self._parts

    @property
    def querystring(self):
        """Yields the URL querystring: "www.example.com?page=1" => "page=1" """
        s = self.parts[QUERY].items()
        s = dict((bytestring(k), v if v is not None else "") for k, v in s)
        if sys.version > "3":
            # Python 3
            s = urlencode(s)
        else:
            # Python 2: urlencode() expects byte strings
            t = {
                key: value.encode("utf-8") if isinstance(value, str) else value
                for key, value in s.items()
            }
            s = urlencode(t).decode("utf-8")
        return s

    def __getattr__(self, k):
        if k in self.__dict__:
            return self.__dict__[k]
        if k in self.parts:
            return self.__dict__["_parts"][k]
        raise AttributeError("'URL' object has no attribute '%s'" % k)

    def __setattr__(self, k, v):
        if k in self.__dict__:
            self.__dict__[k] = u(v)
            return
        if k == "string":
            self._set_string(v)
            return
        if k == "query":
            self.parts[k] = v
            return
        if k in self.parts:
            self.__dict__["_parts"][k] = u(v)
            return
        raise AttributeError("'URL' object has no attribute '%s'" % k)

    def open(
        self,
        timeout=10,
        proxy=None,
        user_agent=USER_AGENT,
        referrer=REFERRER,
        authentication=None,
    ):
        """Returns a connection to the url from which data can be retrieved with connection.read().
        When the timeout amount of seconds is exceeded, raises a URLTimeout.
        When an error occurs, raises a URLError (e.g. HTTP404NotFound).
        """
        url = self.string
        # Handle local files directly
        if os.path.exists(url):
            return urlopen(url)
        # Handle method=POST with query string as a separate parameter.
        post = self.method == POST and self.querystring or None
        socket.setdefaulttimeout(timeout)
        # Handle proxies and cookies.
        handlers = []
        if proxy:
            handlers.append(ProxyHandler({proxy[1]: proxy[0]}))
        handlers.append(HTTPCookieProcessor(cookielib.CookieJar()))
        handlers.append(HTTPHandler)
        install_opener(build_opener(*handlers))
        # Send request.
        try:
            request = Request(
                url, post, {"User-Agent": user_agent, "Referer": referrer}
            )
            # Basic authentication is established with authentication=(username, password).
            if authentication is not None:
                authentication = tuple(encode_utf8(x) for x in authentication)
                request.add_header(
                    "Authorization",
                    "Basic %s"
                    % decode_utf8(base64.b64encode(b"%s:%s" % authentication)),
                )
            return urlopen(request)
        except UrllibHTTPError as e:
            if e.code == 301:
                raise HTTP301Redirect(src=e, url=url)
            if e.code == 400:
                raise HTTP400BadRequest(src=e, url=url)
            if e.code == 401:
                raise HTTP401Authentication(src=e, url=url)
            if e.code == 403:
                raise HTTP403Forbidden(src=e, url=url)
            if e.code == 404:
                raise HTTP404NotFound(src=e, url=url)
            if e.code == 414:
                raise HTTP414RequestURITooLong(src=e, url=url)
            if e.code == 420:
                raise HTTP420Error(src=e, url=url)
            if e.code == 429:
                raise HTTP429TooMayRequests(src=e, url=url)
            if e.code == 500:
                raise HTTP500InternalServerError(src=e, url=url)
            if e.code == 503:
                raise HTTP503ServiceUnavailable(src=e, url=url)
            raise HTTPError(str(e), src=e, url=url)
        except httplib.BadStatusLine as e:
            raise HTTPError(str(e), src=e, url=url)
        except socket.timeout as e:
            raise URLTimeout(src=e, url=url)
        except socket.error as e:
            if "timed out" in str((e.args + ("", ""))[0]) or "timed out" in str(
                (e.args + ("", ""))[1]
            ):
                raise URLTimeout(src=e, url=url)
            raise URLError(str(e), src=e, url=url)
        except UrllibURLError as e:
            if "timed out" in str(e.reason):
                raise URLTimeout(src=e, url=url)
            raise URLError(str(e), src=e, url=url)
        except ValueError as e:
            raise URLError(str(e), src=e, url=url)

    def download(
        self,
        timeout=10,
        cached=True,
        throttle=0,
        proxy=None,
        user_agent=USER_AGENT,
        referrer=REFERRER,
        authentication=None,
        unicode=False,
        **kwargs
    ):
        """Downloads the content at the given URL (by default it will be cached locally).
        Unless unicode=False, the content is returned as a unicode string.
        """
        # Filter OAuth parameters from cache id (they will be unique for each request).
        if self._parts is None and self.method == GET and "oauth_" not in self._string:
            id = self._string
        else:
            id = repr(self.parts)
            id = re.sub("u{0,1}'oauth_.*?': u{0,1}'.*?', ", "", id)
        # Keep a separate cache of unicode and raw download for same URL.
        if unicode is True:
            id = "u" + id
        if cached and id in cache:
            if isinstance(cache, dict):  # Not a Cache object.
                return cache[id]
            if unicode is True:
                return cache[id]
            if unicode is False:
                return cache.get(id, unicode=False)
        t = time.time()
        # Open a connection with the given settings, read it and (by default) cache the data.
        try:
            data = self.open(
                timeout, proxy, user_agent, referrer, authentication
            ).read()
        except socket.timeout as e:
            raise URLTimeout(src=e, url=self.string)
        if unicode is True:
            data = u(data)
        if cached:
            cache[id] = data
        if throttle:
            time.sleep(max(throttle - (time.time() - t), 0))
        return data

    def read(self, *args, **kwargs):
        return self.open(**kwargs).read(*args)

    @property
    def exists(self, timeout=10):
        """Yields False if the URL generates a HTTP404NotFound error."""
        try:
            self.open(timeout)
        except HTTP404NotFound:
            return False
        except HTTPError:
            return True
        except URLTimeout:
            return True
        except URLError:
            return False
        except:
            return True
        return True

    @property
    def mimetype(self, timeout=10):
        """Yields the MIME-type of the document at the URL, or None.
        MIME is more reliable than simply checking the document extension.
        You can then do: URL.mimetype in MIMETYPE_IMAGE.
        """
        try:
            return self.headers["content-type"].split(";")[0]
        except KeyError:
            return None

    @property
    def headers(self, timeout=10):
        """Yields a dictionary with the HTTP response headers."""
        if self.__dict__["_headers"] is None:
            try:
                h = dict(self.open(timeout).info())
            except URLError:
                h = {}
            self.__dict__["_headers"] = h

        # Backward compatibility (Python 2)
        if "Content-Type" in self.__dict__["_headers"]:
            self.__dict__["_headers"]["content-type"] = self.__dict__["_headers"][
                "Content-Type"
            ]

        return self.__dict__["_headers"]

    @property
    def redirect(self, timeout=10):
        """Yields the redirected URL, or None."""
        if self.__dict__["_redirect"] is None:
            try:
                r = u(self.open(timeout).geturl())
            except URLError:
                r = None
            self.__dict__["_redirect"] = r != self.string and r or ""
        return self.__dict__["_redirect"] or None

    def __str__(self):
        # The string representation includes the query attributes with HTTP GET.
        P = self.parts
        u = []
        if P[PROTOCOL]:
            u.append("%s://" % P[PROTOCOL])
        if P[USERNAME]:
            u.append("%s:%s@" % (P[USERNAME], P[PASSWORD]))
        if P[DOMAIN]:
            u.append(P[DOMAIN])
        if P[PORT]:
            u.append(":%s" % P[PORT])
        if P[PORT] or P[DOMAIN] and not P[PATH] and not P[PAGE]:
            u.append("/")
        if P[PATH]:
            u.append("/%s/" % "/".join(P[PATH]))
        if P[PAGE] and len(u) > 0:
            u[-1] = u[-1].rstrip("/")
        if P[PAGE]:
            u.append("/%s" % P[PAGE])
        if P[QUERY] and self.method == GET:
            u.append("?%s" % self.querystring)
        if P[ANCHOR]:
            u.append("#%s" % P[ANCHOR])
        u = "".join(u)
        u = u.lstrip("/")
        return u

    def __repr__(self):
        return "URL(%s, method=%s)" % (repr(self.string), repr(self.method))

    def copy(self):
        return URL(self.string, self.method, self.query)


def download(
    url="",
    method=GET,
    query={},
    timeout=10,
    cached=True,
    throttle=0,
    proxy=None,
    user_agent=USER_AGENT,
    referrer=REFERRER,
    authentication=None,
    unicode=False,
):
    """Downloads the content at the given URL (by default it will be cached locally).
    Unless unicode=False, the content is returned as a unicode string.
    """
    return URL(url, method, query).download(
        timeout, cached, throttle, proxy, user_agent, referrer, authentication, unicode
    )


# url = URL("http://user:pass@example.com:992/animal/bird?species#wings")
# print(url.parts)
# print(url.query)
# print(url.string)

# --- STREAMING URL BUFFER --------------------------------------------------------------------------


def bind(object, method, function):
    """Attaches the function as a method with the given name to the given object."""
    if new:
        # Python 2
        setattr(object, method, new.instancemethod(function, object))
    else:
        # Python 3: There is no good reason to use this function in Python 3.
        setattr(object, method, function)


class Stream(list):
    def __init__(self, url, delimiter="\n", **kwargs):
        """Buffered stream of data from a given URL."""
        self.socket = URL(url).open(**kwargs)
        self.buffer = ""
        self.delimiter = delimiter

    def update(self, bytes=1024):
        """Reads a number of bytes from the stream.
        If a delimiter is encountered, calls Stream.parse() on the packet.
        """
        packets = []
        self.buffer += self.socket.read(bytes).decode("utf-8")
        self.buffer = self.buffer.split(self.delimiter, 1)
        while len(self.buffer) > 1:
            data = self.buffer[0]
            data = self.parse(data)
            if data is not None:
                packets.append(data)
            self.buffer = self.buffer[-1]
            self.buffer = self.buffer.split(self.delimiter, 1)
        self.buffer = self.buffer[-1]
        self.extend(packets)
        return packets

    def parse(self, data):
        """Must be overridden in a subclass."""
        return data

    def clear(self):
        list.__init__(self, [])


def stream(url, delimiter="\n", parse=lambda data: data, **kwargs):
    """Returns a new Stream with the given parse method."""
    stream = Stream(url, delimiter, **kwargs)
    bind(stream, "parse", lambda stream, data: parse(data))
    return stream


# --- FIND URLs -------------------------------------------------------------------------------------
# Functions for parsing URL's and e-mail adresses from strings.

RE_URL_PUNCTUATION = ("\"'{(>", "\"'.,;)}")
RE_URL_HEAD = r"[%s|\[|\s]" % "|".join(
    RE_URL_PUNCTUATION[0]
)  # Preceded by space, parenthesis or HTML tag.
RE_URL_TAIL = r"[%s|\]]*[\s|\<]" % "|".join(
    RE_URL_PUNCTUATION[1]
)  # Followed by space, punctuation or HTML tag.
RE_URL1 = r"(https?://.*?)" + RE_URL_TAIL  # Starts with http:// or https://
RE_URL2 = RE_URL_HEAD + r"(www\..*?\..*?)" + RE_URL_TAIL  # Starts with www.
RE_URL3 = RE_URL_HEAD + r"([\w|-]*?\.(com|net|org|edu|de|uk))" + RE_URL_TAIL

RE_URL1, RE_URL2, RE_URL3 = (
    re.compile(RE_URL1, re.I),
    re.compile(RE_URL2, re.I),
    re.compile(RE_URL3, re.I),
)


def find_urls(string, unique=True):
    """Returns a list of URLs parsed from the string.
    Works on http://, https://, www. links or domain names ending in .com, .org, .net.
    Links can be preceded by leading punctuation (open parens)
    and followed by trailing punctuation (period, comma, close parens).
    """
    string = u(string)
    string = string.replace("\u2024", ".")
    string = string.replace(" ", "  ")
    matches = []
    for p in (RE_URL1, RE_URL2, RE_URL3):
        for m in p.finditer(" %s " % string):
            s = m.group(1)
            s = s.split('">')[0].split("'>")[0]  # google.com">Google => google.com
            if not unique or s not in matches:
                matches.append(s)
    return matches


links = find_urls

RE_EMAIL = re.compile(
    r"[\w\-\.\+]+@(\w[\w\-]+\.)+[\w\-]+"
)  # tom.de+smedt@clips.ua.ac.be


def find_email(string, unique=True):
    """Returns a list of e-mail addresses parsed from the string."""
    string = u(string).replace("\u2024", ".")
    matches = []
    for m in RE_EMAIL.finditer(string):
        s = m.group(0)
        if not unique or s not in matches:
            matches.append(s)
    return matches


def find_between(a, b, string):
    """Returns a list of substrings between a and b in the given string."""
    p = "%s(.*?)%s" % (a, b)
    p = re.compile(p, re.DOTALL | re.I)
    return [m for m in p.findall(string)]


#### PLAIN TEXT ####################################################################################
# Functions for stripping HTML tags from strings.

BLOCK = [
    "title",
    "h1",
    "h2",
    "h3",
    "h4",
    "h5",
    "h6",
    "p",
    "center",
    "blockquote",
    "div",
    "table",
    "ul",
    "ol",
    "dl",
    "pre",
    "code",
    "form",
]

SELF_CLOSING = ["br", "hr", "img"]

# Element tag replacements for a stripped version of HTML source with strip_tags().
# Block-level elements are followed by linebreaks,
# list items are preceded by an asterisk ("*").
LIST_ITEM = "*"
blocks = dict.fromkeys(BLOCK + ["br", "tr", "td"], ("", "\n\n"))
blocks.update(
    {
        "li": ("%s " % LIST_ITEM, "\n"),
        "img": ("", ""),
        "br": ("", "\n"),
        "th": ("", "\n"),
        "tr": ("", "\n"),
        "td": ("", "\t"),
    }
)


class HTMLParser(_HTMLParser):
    def clean(self, html):
        html = decode_utf8(html)
        html = html.replace("/>", " />")
        html = html.replace("  />", " />")
        html = html.replace("<!", "&lt;!")
        html = html.replace("&lt;!DOCTYPE", "<!DOCTYPE")
        html = html.replace("&lt;!doctype", "<!doctype")
        html = html.replace("&lt;!--", "<!--")
        return html


class HTMLTagstripper(HTMLParser):
    def __init__(self):
        HTMLParser.__init__(self)

    def strip(self, html, exclude=[], replace=blocks):
        """Returns the HTML string with all element tags (e.g. <p>) removed.
        - exclude    : a list of tags to keep. Element attributes are stripped.
                       To preserve attributes a dict of (tag name, [attribute])-items can be given.
        - replace    : a dictionary of (tag name, (replace_before, replace_after))-items.
                       By default, block-level elements are separated with linebreaks.
        """
        if html is None:
            return None
        self._exclude = (
            isinstance(exclude, dict) and exclude or dict.fromkeys(exclude, [])
        )
        self._replace = replace
        self._data = []
        self.feed(self.clean(html))
        self.close()
        self.reset()
        return "".join(self._data)

    def handle_starttag(self, tag, attributes):
        if tag in BLOCK and self._data and self._data[-1][-1:] != "\n":
            # Block-level elements always break to a new line.
            self._data.append("\n")
        if tag in self._exclude:
            # Create the tag attribute string,
            # including attributes defined in the HTMLTagStripper._exclude dict.
            a = len(self._exclude[tag]) > 0 and attributes or []
            a = ['%s="%s"' % (k, v) for k, v in a if k in self._exclude[tag]]
            a = (" " + " ".join(a)).rstrip()
            self._data.append("<%s%s>" % (tag, a))
        if tag in self._replace:
            self._data.append(self._replace[tag][0])
        if tag in self._replace and tag in SELF_CLOSING:
            self._data.append(self._replace[tag][1])

    def handle_endtag(self, tag):
        if tag in self._exclude and self._data and self._data[-1].startswith("<" + tag):
            # Never keep empty elements (e.g. <a></a>).
            self._data.pop(-1)
            return
        if tag in self._exclude:
            self._data.append("</%s>" % tag)
        if tag in self._replace:
            self._data.append(self._replace[tag][1])

    def handle_data(self, data):
        self._data.append(data.strip("\n\t"))

    def handle_comment(self, comment):
        if "comment" in self._exclude or "!--" in self._exclude:
            self._data.append("<!--%s-->" % comment)


# As a function:
strip_tags = HTMLTagstripper().strip


def strip_element(string, tag, attributes=""):
    """Removes all elements with the given tagname and attributes from the string.
    Open and close tags are kept in balance.
    No HTML parser is used: strip_element(s, "a", 'class="x"') matches
    '<a class="x">' or '<a href="x" class="x">' but not "<a class='x'>".
    """
    s = string.lower()  # Case-insensitive.
    t = tag.strip("</>")
    a = (" " + attributes.lower().strip()).rstrip()
    i = 0
    j = 0
    while j >= 0:
        # i = s.find("<%s%s" % (t, a), i)
        m = re.search(r"<%s[^\>]*?%s" % (t, a), s[i:])
        i = i + m.start() if m else -1
        j = s.find("</%s>" % t, i + 1)
        opened, closed = s[i:j].count("<%s" % t), 1
        while opened > closed and j >= 0:
            k = s.find("</%s>" % t, j + 1)
            opened += s[j:k].count("<%s" % t)
            closed += 1
            j = k
        if i < 0:
            return string
        if j < 0:
            return string[:i]
        string = string[:i] + string[j + len(t) + 3 :]
        s = string.lower()
    return string


def strip_between(a, b, string):
    """Removes anything between (and including) string a and b inside the given string."""
    p = "%s.*?%s" % (a, b)
    p = re.compile(p, re.DOTALL | re.I)
    return re.sub(p, "", string)


def strip_javascript(html):
    return strip_between("<script.*?>", "</script>", html)


def strip_inline_css(html):
    return strip_between("<style.*?>", "</style>", html)


def strip_comments(html):
    return strip_between("<!--", "-->", html)


def strip_forms(html):
    return strip_between("<form.*?>", "</form>", html)


RE_AMPERSAND = re.compile("\&(?!\#)")  # & not followed by #
RE_UNICODE = re.compile(r"&(#?)(x|X?)(\w+);")  # &#201;


def encode_entities(string):
    """Encodes HTML entities in the given string ("<" => "&lt;").
    For example, to display "<em>hello</em>" in a browser,
    we need to pass "&lt;em&gt;hello&lt;/em&gt;" (otherwise "hello" in italic is displayed).
    """
    if isinstance(string, str):
        string = RE_AMPERSAND.sub("&amp;", string)
        string = string.replace("<", "&lt;")
        string = string.replace(">", "&gt;")
        string = string.replace('"', "&quot;")
        string = string.replace("'", "&#39;")
    return string


def decode_entities(string):
    """Decodes HTML entities in the given string ("&lt;" => "<")."""
    # http://snippets.dzone.com/posts/show/4569
    def replace_entity(match):
        hash, hex, name = match.group(1), match.group(2), match.group(3)
        if hash == "#" or name.isdigit():
            if hex == "":
                return chr(int(name))  # "&#38;" => "&"
            if hex.lower() == "x":
                return chr(int("0x" + name, 16))  # "&#x0026;" = > "&"
        else:
            cp = name2codepoint.get(name)  # "&amp;" => "&"
            return chr(cp) if cp else match.group()  # "&foo;" => "&foo;"

    if isinstance(string, str):
        return RE_UNICODE.subn(replace_entity, string)[0]
    return string


def encode_url(string):
    return quote_plus(bytestring(string))  # "black/white" => "black%2Fwhite".


def decode_url(string):
    return u(unquote_plus(bytestring(string)))


RE_SPACES = re.compile("( |\xa0)+", re.M)  # Matches one or more spaces.
RE_TABS = re.compile(r"\t+", re.M)  # Matches one or more tabs.


def collapse_spaces(string, indentation=False, replace=" "):
    """Returns a string with consecutive spaces collapsed to a single space.
    Whitespace on empty lines and at the end of each line is removed.
    With indentation=True, retains leading whitespace on each line.
    """
    p = []
    for x in string.splitlines():
        n = indentation and len(x) - len(x.lstrip()) or 0
        p.append(x[:n] + RE_SPACES.sub(replace, x[n:]).strip())
    return "\n".join(p)


def collapse_tabs(string, indentation=False, replace=" "):
    """Returns a string with (consecutive) tabs replaced by a single space.
    Whitespace on empty lines and at the end of each line is removed.
    With indentation=True, retains leading whitespace on each line.
    """
    p = []
    for x in string.splitlines():
        n = indentation and len(x) - len(x.lstrip()) or 0
        p.append(x[:n] + RE_TABS.sub(replace, x[n:]).strip())
    return "\n".join(p)


def collapse_linebreaks(string, threshold=1):
    """Returns a string with consecutive linebreaks collapsed to at most the given threshold.
    Whitespace on empty lines and at the end of each line is removed.
    """
    n = "\n" * threshold
    p = [s.rstrip() for s in string.splitlines()]
    string = "\n".join(p)
    string = re.sub(n + r"+", n, string)
    return string


def plaintext(html, keep=[], replace=blocks, linebreaks=2, indentation=False):
    """Returns a string with all HTML tags removed.
    Content inside HTML comments, the <style> tag and the <script> tags is removed.
    - keep        : a list of tags to keep. Element attributes are stripped.
                    To preserve attributes a dict of (tag name, [attribute])-items can be given.
    - replace     : a dictionary of (tag name, (replace_before, replace_after))-items.
                    By default, block-level elements are followed by linebreaks.
    - linebreaks  : the maximum amount of consecutive linebreaks,
    - indentation : keep left line indentation (tabs and spaces)?
    """
    if isinstance(html, Element):
        html = html.content
    if not keep.__contains__("script"):
        html = strip_javascript(html)
    if not keep.__contains__("style"):
        html = strip_inline_css(html)
    if not keep.__contains__("form"):
        html = strip_forms(html)
    if not keep.__contains__("comment") and not keep.__contains__("!--"):
        html = strip_comments(html)
    html = html.replace("\r", "\n")
    html = decode_entities(html)
    html = strip_tags(html, exclude=keep, replace=replace)
    html = collapse_spaces(html, indentation)
    html = collapse_tabs(html, indentation)
    html = collapse_linebreaks(html, linebreaks)
    html = html.strip()
    return html


#### SEARCH ENGINE #################################################################################

SEARCH = "search"  # Query for pages (i.e. links to websites).
IMAGE = "image"  # Query for images.
NEWS = "news"  # Query for news items.

TINY = "tiny"  # Image size around 100x100.
SMALL = "small"  # Image size around 200x200.
MEDIUM = "medium"  # Image size around 500x500.
LARGE = "large"  # Image size around 1000x1000.

RELEVANCY = "relevancy"  # Sort results by most relevant.
LATEST = "latest"  # Sort results by most recent.


class Result(dict):
    def __init__(self, url, **kwargs):
        """An item in a list of results returned by SearchEngine.search().
        All dictionary keys are available as Unicode string attributes.
        - id       : unique identifier,
        - url      : the URL of the referred web content,
        - title    : the title of the content at the URL,
        - text     : the content text,
        - language : the content language,
        - author   : for news items and posts, the author,
        - date     : for news items and posts, the publication date.
        """
        dict.__init__(self)
        self.url = url
        self.id = kwargs.pop("id", "")
        self.title = kwargs.pop("title", "")
        self.text = kwargs.pop("text", "")
        self.language = kwargs.pop("language", "")
        self.author = kwargs.pop("author", "")
        self.date = kwargs.pop("date", "")
        self.votes = kwargs.pop("votes", 0)  # (e.g., Facebook likes)
        self.shares = kwargs.pop("shares", 0)  # (e.g., Twitter retweets)
        self.comments = kwargs.pop("comments", 0)
        for k, v in kwargs.items():
            self[k] = v

    @property
    def txt(self):
        return self.text

    @property
    def description(self):
        return self.text  # Backwards compatibility.

    @property
    def likes(self):
        return self.votes

    @property
    def retweets(self):
        return self.shares

    def download(self, *args, **kwargs):
        """Download the content at the given URL.
        By default it will be cached - see URL.download().
        """
        return URL(self.url).download(*args, **kwargs)

    def _format(self, v):
        if isinstance(v, bytes):  # Store strings as unicode.
            return u(v)
        if v is None:
            return ""
        return v

    def __getattr__(self, k):
        return self.get(k, "")

    def __getitem__(self, k):
        return self.get(k, "")

    def __setattr__(self, k, v):
        self.__setitem__(k, v)

    def __setitem__(self, k, v):
        dict.__setitem__(self, u(k), self._format(v))

    def setdefault(self, k, v=None):
        return dict.setdefault(self, u(k), self._format(v))

    def update(self, *args, **kwargs):
        dict.update(
            self, [(u(k), self._format(v)) for k, v in dict(*args, **kwargs).items()]
        )

    def __repr__(self):
        return "Result(%s)" % repr(dict((k, v) for k, v in self.items() if v))


class Results(list):
    def __init__(self, source=None, query=None, type=SEARCH, total=0):
        """A list of results returned from SearchEngine.search().
        - source: the service that yields the results (e.g. GOOGLE, TWITTER).
        - query : the query that yields the results.
        - type  : the query type (SEARCH, IMAGE, NEWS).
        - total : the total result count.
                  This is not the length of the list, but the total number of matches for the given query.
        """
        self.source = source
        self.query = query
        self.type = type
        self.total = total


class SearchEngine(object):
    def __init__(self, license=None, throttle=1.0, language=None):
        """A base class for a web service.
        - license  : license key for the API,
        - throttle : delay between requests (avoid hammering the server).
        Inherited by: Google, Bing, Wikipedia, Twitter, Facebook, Flickr, ...
        """
        self.license = license
        self.throttle = throttle  # Amount of sleep time after executing a query.
        self.language = language  # Result.language restriction (e.g., "en").
        self.format = lambda x: x  # Formatter applied to each attribute of each Result.

    def search(
        self,
        query,
        type=SEARCH,
        start=1,
        count=10,
        sort=RELEVANCY,
        size=None,
        cached=True,
        **kwargs
    ):
        return Results(source=None, query=query, type=type)


class SearchEngineError(HTTPError):
    pass


class SearchEngineTypeError(SearchEngineError):
    pass  # Raised when an unknown type is passed to SearchEngine.search().


class SearchEngineLimitError(SearchEngineError):
    pass  # Raised when the query limit for a license is reached.


# print(sort(["black", "happy"], "darth vader", GOOGLE))

#### DOCUMENT OBJECT MODEL #########################################################################
# The Document Object Model (DOM) is a cross-platform and language-independent convention
# for representing and interacting with objects in HTML, XHTML and XML documents.
# The pattern.web DOM can be used to traverse HTML source code as a tree of nested elements.
# The pattern.web DOM is based on Beautiful Soup.

# Beautiful Soup is wrapped in DOM, Element and Text classes, resembling the Javascript DOM.
# Beautiful Soup can also be used directly, since it is imported here.
# L. Richardson (2004), http://www.crummy.com/software/BeautifulSoup/

SOUP = (
    BeautifulSoup.BeautifulSoup,
    BeautifulSoup.Tag,
    BeautifulSoup.NavigableString,
    BeautifulSoup.Comment,
)

NODE, TEXT, COMMENT, ELEMENT, DOCUMENT = (
    "node",
    "text",
    "comment",
    "element",
    "document",
)

# --- NODE ------------------------------------------------------------------------------------------


class Node(object):
    def __init__(self, html, type=NODE, **kwargs):
        """The base class for Text, Comment and Element.
        All DOM nodes can be navigated in the same way (e.g. Node.parent, Node.children, ...)
        """
        self.type = type
        self._p = (
            not isinstance(html, SOUP)
            and BeautifulSoup.BeautifulSoup(u(html), "lxml", **kwargs)
            or html
        )

    @property
    def _beautifulSoup(self):
        # If you must, access the BeautifulSoup object with Node._beautifulSoup.
        return self._p

    def __eq__(self, other):
        # Two Node objects containing the same BeautifulSoup object, are the same.
        return isinstance(other, Node) and hash(self._p) == hash(other._p)

    def _wrap(self, x):
        # Navigating to other nodes yields either Text, Element or None.
        if isinstance(x, BeautifulSoup.Comment):
            return Comment(x)
        if isinstance(x, BeautifulSoup.Declaration):
            return Text(x)
        if isinstance(x, BeautifulSoup.NavigableString):
            return Text(x)
        if isinstance(x, BeautifulSoup.Tag):
            return Element(x)

    @property
    def parent(self):
        return self._wrap(self._p.parent)

    @property
    def children(self):
        return (
            hasattr(self._p, "contents")
            and [self._wrap(x) for x in self._p.contents]
            or []
        )

    @property
    def html(self):
        return self.__str__()

    @property
    def source(self):
        return self.__str__()

    @property
    def next_sibling(self):
        return self._wrap(self._p.next_sibling)

    @property
    def previous_sibling(self):
        return self._wrap(self._p.previous_sibling)

    next, prev, previous = next_sibling, previous_sibling, previous_sibling

    def traverse(self, visit=lambda node: None):
        """Executes the visit function on this node and each of its child nodes."""
        visit(self)
        [node.traverse(visit) for node in self.children]

    def remove(self, child):
        """Removes the given child node (and all nested nodes)."""
        child._p.extract()

    def __bool__(self):
        return True

    def __len__(self):
        return len(self.children)

    def __iter__(self):
        return iter(self.children)

    def __getitem__(self, index):
        return self.children[index]

    def __repr__(self):
        return "Node(type=%s)" % repr(self.type)

    def __str__(self):
        return u(self._p)

    def __call__(self, *args, **kwargs):
        pass


# --- TEXT ------------------------------------------------------------------------------------------


class Text(Node):
    """Text represents a chunk of text without formatting in a HTML document.
    For example: "the <b>cat</b>" is parsed to [Text("the"), Element("cat")].
    """

    def __init__(self, string):
        Node.__init__(self, string, type=TEXT)

    def __repr__(self):
        return "Text(%s)" % repr(self._p)


class Comment(Text):
    """Comment represents a comment in the HTML source code.
    For example: "<!-- comment -->".
    """

    def __init__(self, string):
        Node.__init__(self, string, type=COMMENT)

    def __repr__(self):
        return "Comment(%s)" % repr(self._p)


# --- ELEMENT ---------------------------------------------------------------------------------------


class Element(Node):
    def __init__(self, html):
        """Element represents an element or tag in the HTML source code.
        For example: "<b>hello</b>" is a "b"-Element containing a child Text("hello").
        """
        Node.__init__(self, html, type=ELEMENT)

    @property
    def tagname(self):
        return self._p.name

    tag = tagName = tagname

    @property
    def attributes(self):
        return self._p.attrs

    attr = attrs = attributes

    @property
    def id(self):
        return self.attributes.get("id")

    @property
    def content(self):
        """Yields the element content as a unicode string."""
        return "".join([u(x) for x in self._p.contents])

    string = content

    @property
    def source(self):
        """Yields the HTML source as a unicode string (tag + content)."""
        return u(self._p)

    html = src = source

    def get_elements_by_tagname(self, v):
        """Returns a list of nested Elements with the given tag name.
        The tag name can include a class (e.g. div.header) or an id (e.g. div#content).
        """
        if isinstance(v, str) and "#" in v:
            v1, v2 = v.split("#")
            v1 = v1 in ("*", "") or v1.lower()
            return [Element(x) for x in self._p.find_all(v1, id=v2)]
        if isinstance(v, str) and "." in v:
            v1, v2 = v.split(".")
            v1 = v1 in ("*", "") or v1.lower()
            return [Element(x) for x in self._p.find_all(v1, v2)]
        return [Element(x) for x in self._p.find_all(v in ("*", "") or v.lower())]

    by_tag = getElementsByTagname = get_elements_by_tagname

    def get_element_by_id(self, v):
        """Returns the first nested Element with the given id attribute value."""
        return ([Element(x) for x in self._p.find_all(id=v, limit=1) or []] + [None])[0]

    by_id = getElementById = get_element_by_id

    def get_elements_by_classname(self, v):
        """Returns a list of nested Elements with the given class attribute value."""
        return [Element(x) for x in (self._p.find_all(True, v))]

    by_class = getElementsByClassname = get_elements_by_classname

    def get_elements_by_attribute(self, **kwargs):
        """Returns a list of nested Elements with the given attribute value."""
        return [Element(x) for x in (self._p.find_all(True, attrs=kwargs))]

    by_attribute = by_attr = getElementsByAttribute = get_elements_by_attribute

    def __call__(self, selector):
        """Returns a list of nested Elements that match the given CSS selector.
        For example: Element("div#main p.comment a:first-child") matches:
        """
        return SelectorChain(selector).search(self)

    def __getattr__(self, k):
        if k in self.__dict__:
            return self.__dict__[k]
        if k in self.attributes:
            return self.attributes[k]
        raise AttributeError("'Element' object has no attribute '%s'" % k)

    def __contains__(self, v):
        if isinstance(v, Element):
            v = v.content
        return v in self.content

    def __repr__(self):
        return "Element(tag=%s)" % repr(self.tagname)


# --- DOCUMENT --------------------------------------------------------------------------------------


class Document(Element):
    def __init__(self, html, **kwargs):
        """Document is the top-level element in the Document Object Model.
        It contains nested Element, Text and Comment nodes.
        """
        # Aliases for BeautifulSoup optional parameters:
        # kwargs["selfClosingTags"] = kwargs.pop("self_closing", kwargs.get("selfClosingTags"))
        Node.__init__(self, u(html).strip(), type=DOCUMENT, **kwargs)

    @property
    def declaration(self):
        """Yields the <!doctype> declaration, as a TEXT Node or None."""
        for child in self.children:
            if isinstance(child._p, (BeautifulSoup.Declaration, BeautifulSoup.Doctype)):
                return child

    @property
    def head(self):
        return self._wrap(self._p.head)

    @property
    def body(self):
        return self._wrap(self._p.body)

    @property
    def tagname(self):
        return None

    tag = tagname

    def __repr__(self):
        return "Document()"


DOM = Document

# article = Wikipedia().search("Document Object Model")
# dom = DOM(article.html)
# print(dom.get_element_by_id("References").source)
# print([element.attributes["href"] for element in dom.get_elements_by_tagname("a")])
# print(dom.get_elements_by_tagname("p")[0].next.previous.children[0].parent.__class__)
# print()

# --- DOM CSS SELECTORS -----------------------------------------------------------------------------
# CSS selectors are pattern matching rules (or selectors) to select elements in the DOM.
# CSS selectors may range from simple element tag names to rich contextual patterns.
# http://www.w3.org/TR/CSS2/selector.html

# "*"                 =  <div>, <p>, ...                (all elements)
# "*#x"               =  <div id="x">, <p id="x">, ...  (all elements with id="x")
# "div#x"             =  <div id="x">                   (<div> elements with id="x")
# "div.x"             =  <div class="x">                (<div> elements with class="x")
# "div[class='x']"    =  <div class="x">                (<div> elements with attribute "class"="x")
# "div:contains('x')" =  <div>xyz</div>                 (<div> elements that contain "x")
# "div:first-child"   =  <div><a>1st<a><a></a></div>    (first child inside a <div>)
# "div a"             =  <div><p><a></a></p></div>      (all <a>'s inside a <div>)
# "div, a"            =  <div>, <a>                     (all <a>'s and <div> elements)
# "div + a"           =  <div></div><a></a>             (all <a>'s directly preceded by <div>)
# "div > a"           =  <div><a></a></div>             (all <a>'s directly inside a <div>)
# "div < a"                                            (all <div>'s directly containing an <a>)

# Selectors are case-insensitive.


def _encode_space(s):
    return s.replace(" ", "<!space!>")


def _decode_space(s):
    return s.replace("<!space!>", " ")


class Selector(object):
    def __init__(self, s):
        """A simple CSS selector is a type (e.g., "p") or universal ("*") selector
        followed by id selectors, attribute selectors, or pseudo-elements.
        """
        self.string = s
        s = s.strip()
        s = s.lower()
        s = s.startswith(("#", ".", ":")) and "*" + s or s
        s = s.replace("#", " #") + " #"  # #id
        s = s.replace(".", " .")  # .class
        s = s.replace(":", " :")  # :pseudo-element
        s = s.replace("[", " [")  # [attribute="value"]
        s = re.sub(r"\[.*?\]", lambda m: re.sub(r" (\#|\.|\:)", "\\1", m.group(0)), s)
        s = re.sub(r"\[.*?\]", lambda m: _encode_space(m.group(0)), s)
        s = re.sub(r":contains\(.*?\)", lambda m: _encode_space(m.group(0)), s)
        s = s.split(" ")
        self.tag, self.id, self.classes, self.pseudo, self.attributes = (
            s[0],
            [x[1:] for x in s if x[0] == "#"][0],
            set([x[1:] for x in s if x[0] == "."]),
            set([x[1:] for x in s if x[0] == ":"]),
            dict(self._parse_attribute(x) for x in s if x[0] == "["),
        )

    def _parse_attribute(self, s):
        """Returns an (attribute, value)-tuple for the given attribute selector."""
        s = s.strip("[]")
        s = s.replace("'", "")
        s = s.replace('"', "")
        s = _decode_space(s)
        s = re.sub(r"(\~|\||\^|\$|\*)\=", "=\\1", s)
        s = s.split("=") + [True]
        s = s[:2]
        if s[1] is not True:
            r = r"^%s$"
            if s[1].startswith(("~", "|", "^", "$", "*")):
                p, s[1] = s[1][0], s[1][1:]
                if p == "~":
                    r = r"(^|\s)%s(\s|$)"
                if p == "|":
                    r = r"^%s(-|$)"  # XXX doesn't work with spaces.
                if p == "^":
                    r = r"^%s"
                if p == "$":
                    r = r"%s$"
                if p == "*":
                    r = r"%s"
            s[1] = re.compile(r % s[1], re.I)
        return s[:2]

    def _first_child(self, e):
        """Returns the first child Element of the given element."""
        if isinstance(e, Node):
            for e in e.children:
                if isinstance(e, Element):
                    return e

    def _next_sibling(self, e):
        """Returns the first next sibling Element of the given element."""
        while isinstance(e, Node):
            e = e.next_element
            if isinstance(e, Element):
                return e

    def _previous_sibling(self, e):
        """Returns the last previous sibling Element of the given element."""
        while isinstance(e, Node):
            e = e.previous_element
            if isinstance(e, Element):
                return e

    def _contains(self, e, s):
        """Returns True if string s occurs in the given element (case-insensitive)."""
        s = re.sub(r"^contains\((.*?)\)$", "\\1", s)
        s = re.sub(r"^[\"']|[\"']$", "", s)
        s = _decode_space(s)
        return re.search(s.lower(), e.content.lower()) is not None

    def match(self, e):
        """Returns True if the given element matches the simple CSS selector."""
        if not isinstance(e, Element):
            return False
        if self.tag not in (e.tag, "*"):
            return False
        if self.id not in ((e.id or "").lower(), "", None):
            return False
        if (
            self.classes.issubset(
                set(map(lambda s: s.lower(), e.attr.get("class", "")))
            )
            is False
        ):
            return False
        if "first-child" in self.pseudo and self._first_child(e.parent) != e:
            return False
        if any(
            x.startswith("contains") and not self._contains(e, x) for x in self.pseudo
        ):
            return False  # jQuery :contains("...") selector.
        for k, v in self.attributes.items():
            if k not in e.attrs or not (
                v is True or re.search(v, " ".join(e.attrs[k])) is not None
            ):
                return False
        return True

    def search(self, e):
        """Returns the nested elements that match the simple CSS selector."""
        # Map tag to True if it is "*".
        tag = self.tag == "*" or self.tag
        # Map id into a case-insensitive **kwargs dict.
        i = lambda s: re.compile(r"\b%s(?=$|\s)" % s, re.I)
        a = {"id": i(self.id)} if self.id else {}
        a.update(list(map(lambda kv: (kv[0], kv[1]), list(self.attributes.items()))))
        # Match tag + id + all classes + relevant pseudo-elements.
        if not isinstance(e, Element):
            return []
        if len(self.classes) == 0 or len(self.classes) >= 2:
            e = list(map(Element, e._p.find_all(tag, attrs=a)))
        if len(self.classes) == 1:
            e = list(
                map(
                    Element,
                    e._p.find_all(
                        tag, attrs=dict(a, **{"class": i(list(self.classes)[0])})
                    ),
                )
            )
        if len(self.classes) >= 2:
            e = list(
                filter(lambda e: self.classes.issubset(set(e.attr.get("class", ""))), e)
            )
        if "first-child" in self.pseudo:
            e = list(filter(lambda e: e == self._first_child(e.parent), e))
        if any(x.startswith("contains") for x in self.pseudo):
            e = list(
                filter(
                    lambda e: all(
                        not x.startswith("contains") or self._contains(e, x)
                        for x in self.pseudo
                    ),
                    e,
                )
            )
        return e

    def __repr__(self):
        return "Selector(%s)" % repr(self.string)


class SelectorChain(list):
    def __init__(self, s):
        """A selector is a chain of one or more simple selectors,
        separated by combinators (e.g., ">").
        """
        self.string = s
        for s in s.split(","):
            s = s.lower()
            s = s.strip()
            s = re.sub(r" +", " ", s)
            s = re.sub(r" *\> *", " >", s)
            s = re.sub(r" *\< *", " <", s)
            s = re.sub(r" *\+ *", " +", s)
            s = re.sub(r"\[.*?\]", lambda m: _encode_space(m.group(0)), s)
            s = re.sub(r":contains\(.*?\)", lambda m: _encode_space(m.group(0)), s)
            self.append([])
            for s in s.split(" "):
                if not s.startswith((">", "<", "+")):
                    self[-1].append((" ", Selector(s)))
                elif s.startswith(">"):
                    self[-1].append((">", Selector(s[1:])))
                elif s.startswith("<"):
                    self[-1].append(("<", Selector(s[1:])))
                elif s.startswith("+"):
                    self[-1].append(("+", Selector(s[1:])))

    def search(self, e):
        """Returns the nested elements that match the CSS selector chain."""
        m, root = [], e
        for chain in self:
            e = [root]
            for combinator, s in chain:
                # Search Y, where:
                if combinator == " ":
                    # X Y => X is ancestor of Y
                    e = list(map(s.search, e))
                    e = list(itertools.chain(*e))
                if combinator == ">":
                    # X > Y => X is parent of Y
                    e = list(map(lambda e: list(filter(s.match, e.children)), e))
                    e = list(itertools.chain(*e))
                if combinator == "<":
                    # X < Y => X is child of Y
                    e = list(map(lambda e: e.parent, e))
                    e = list(filter(s.match, e))
                if combinator == "+":
                    # X + Y => X directly precedes Y
                    e = list(map(s._next_sibling, e))
                    e = list(filter(s.match, e))
            m.extend(e)
        return m


# dom = DOM("""
# <html>
# <head></head>
# <body>
#    <div id="#main">
#        <span class="11 22 33">x</span>
#    </div>
# </body>
# </hmtl>
# """)
#
# print(dom("*[class='11']"))
# print(dom("*[class^='11']"))
# print(dom("*[class~='22']"))
# print(dom("*[class$='33']"))
# print(dom("*[class*='3']"))

#### WEB CRAWLER ###################################################################################
# Tested with a crawl across 1,000 domains so far.


class Link(object):
    def __init__(self, url, text="", relation="", referrer=""):
        """A hyperlink parsed from a HTML document, in the form:
        <a href="url"", title="text", rel="relation">xxx</a>.
        """
        self.url, self.text, self.relation, self.referrer = (
            u(url),
            u(text),
            u(relation),
            u(referrer),
        )

    @property
    def description(self):
        return self.text

    def __repr__(self):
        return "Link(url=%s)" % repr(self.url)

    # Used for sorting in Crawler.links:
    def __eq__(self, link):
        return self.url == link.url

    def __ne__(self, link):
        return self.url != link.url

    def __lt__(self, link):
        return self.url < link.url

    def __gt__(self, link):
        return self.url > link.url


class HTMLLinkParser(HTMLParser):
    def __init__(self):
        HTMLParser.__init__(self)

    def parse(self, html, url=""):
        """Returns a list of Links parsed from the given HTML string."""
        if html is None:
            return None
        self._url = url
        self._data = []
        self.feed(self.clean(html))
        self.close()
        self.reset()
        return self._data

    def handle_starttag(self, tag, attributes):
        if tag == "a":
            attributes = dict(attributes)
            if "href" in attributes:
                link = Link(
                    url=attributes.get("href"),
                    text=attributes.get("title"),
                    relation=attributes.get("rel", ""),
                    referrer=self._url,
                )
                self._data.append(link)


def base(url):
    """Returns the URL domain name:
    http://en.wikipedia.org/wiki/Web_crawler => en.wikipedia.org
    """
    return urlparse(url).netloc


def abs(url, base=None):
    """Returns the absolute URL:
    ../media + http://en.wikipedia.org/wiki/ => http://en.wikipedia.org/media
    """
    if url.startswith("#") and base is not None and not base.endswith("/"):
        if not re.search("[^/]/[^/]", base):
            base += "/"
    return urljoin(base, url)


DEPTH = "depth"
BREADTH = "breadth"

FIFO = "fifo"  # First In, First Out.
FILO = "filo"  # First In, Last Out.
LIFO = "lifo"  # Last In, First Out (= FILO).


class Crawler(object):
    def __init__(
        self, links=[], domains=[], delay=20.0, parse=HTMLLinkParser().parse, sort=FIFO
    ):
        """A crawler can be used to browse the web in an automated manner.
        It visits the list of starting URLs, parses links from their content, visits those, etc.
        - Links can be prioritized by overriding Crawler.priority().
        - Links can be ignored by overriding Crawler.follow().
        - Each visited link is passed to Crawler.visit(), which can be overridden.
        """
        self.parse = parse
        self.delay = delay  # Delay between visits to the same (sub)domain.
        self.domains = domains  # Domains the crawler is allowed to visit.
        self.history = {}  # Domain name => time last visited.
        self.visited = {}  # URLs visited.
        self._queue = []  # URLs scheduled for a visit: (priority, time, Link).
        self._queued = {}  # URLs scheduled so far, lookup dictionary.
        self.QUEUE = 10000  # Increase or decrease according to available memory.
        self.sort = sort
        # Queue given links in given order:
        for link in isinstance(links, str) and [links] or links:
            self.push(link, priority=1.0, sort=FIFO)

    @property
    def done(self):
        """Yields True if no further links are scheduled to visit."""
        return len(self._queue) == 0

    def push(self, link, priority=1.0, sort=FILO):
        """Pushes the given link to the queue.
        Position in the queue is determined by priority.
        Equal ranks are sorted FIFO or FILO.
        With priority=1.0 and FILO, the link is inserted to the queue.
        With priority=0.0 and FIFO, the link is appended to the queue.
        """
        if not isinstance(link, Link):
            link = Link(url=link)
        dt = time.time()
        dt = sort == FIFO and dt or 1 / dt
        bisect.insort(self._queue, (1 - priority, dt, link))
        self._queued[link.url] = True

    def pop(self, remove=True):
        """Returns the next Link queued to visit and removes it from the queue.
        Links on a recently visited (sub)domain are skipped until Crawler.delay has elapsed.
        """
        now = time.time()
        for i, (priority, dt, link) in enumerate(self._queue):
            if self.delay <= now - self.history.get(base(link.url), 0):
                if remove is True:
                    self._queue.pop(i)
                    self._queued.pop(link.url, None)
                return link

    @property
    def next(self):
        """Returns the next Link queued to visit (without removing it)."""
        return self.pop(remove=False)

    def crawl(self, method=DEPTH, **kwargs):
        """Visits the next link in Crawler._queue.
        If the link is on a domain recently visited (< Crawler.delay) it is skipped.
        Parses the content at the link for new links and adds them to the queue,
        according to their Crawler.priority().
        Visited links (and content) are passed to Crawler.visit().
        """
        link = self.pop()
        if link is None:
            return False
        if link.url not in self.visited:
            t = time.time()
            url = URL(link.url)
            if url.mimetype == "text/html":
                try:
                    kwargs.setdefault("unicode", True)
                    html = url.download(**kwargs)
                    for new in self.parse(html, url=link.url):
                        new.url = abs(new.url, base=url.redirect or link.url)
                        new.url = self.normalize(new.url)
                        # 1) Parse new links from HTML web pages.
                        # 2) Schedule unknown links for a visit.
                        # 3) Only links that are not already queued are queued.
                        # 4) Only links for which Crawler.follow() is True are queued.
                        # 5) Only links on Crawler.domains are queued.
                        if new.url == link.url:
                            continue
                        if new.url in self.visited:
                            continue
                        if new.url in self._queued:
                            continue
                        if self.follow(new) is False:
                            continue
                        if self.domains and not base(new.url).endswith(
                            tuple(self.domains)
                        ):
                            continue
                        # 6) Limit the queue (remove tail), unless you are Google.
                        if self.QUEUE is not None and self.QUEUE * 1.25 < len(
                            self._queue
                        ):
                            self._queue = self._queue[: self.QUEUE]
                            self._queued.clear()
                            self._queued.update(
                                dict((q[2].url, True) for q in self._queue)
                            )
                        # 7) Position in the queue is determined by Crawler.priority().
                        # 8) Equal ranks are sorted FIFO or FILO.
                        self.push(
                            new,
                            priority=self.priority(new, method=method),
                            sort=self.sort,
                        )
                    self.visit(link, source=html)
                except URLError:
                    # URL can not be reached (HTTP404NotFound, URLTimeout).
                    self.fail(link)
            else:
                # URL MIME-type is not HTML, don't know how to handle.
                self.fail(link)
            # Log the current time visited for the domain (see Crawler.pop()).
            # Log the URL as visited.
            self.history[base(link.url)] = t
            self.visited[link.url] = True
            return True
        # Nothing happened, we already visited this link.
        return False

    def normalize(self, url):
        """Called from Crawler.crawl() to normalize URLs.
        For example: return url.split("?")[0]
        """
        # All links pass through here (visited or not).
        # This can be a place to count backlinks.
        return url

    def follow(self, link):
        """Called from Crawler.crawl() to determine if it should follow this link.
        For example: return "nofollow" not in link.relation
        """
        return True

    def priority(self, link, method=DEPTH):
        """Called from Crawler.crawl() to determine the priority of this link,
        as a number between 0.0-1.0. Links with higher priority are visited first.
        """
        # Depth-first search dislikes external links to other (sub)domains.
        external = base(link.url) != base(link.referrer)
        if external is True:
            if method == DEPTH:
                return 0.75
            if method == BREADTH:
                return 0.85
        return 0.80

    def visit(self, link, source=None):
        """Called from Crawler.crawl() when the link is crawled.
        When source=None, the link is not a web page (and was not parsed),
        or possibly a URLTimeout occured (content size too big).
        """

    def fail(self, link):
        """Called from Crawler.crawl() for link whose MIME-type could not be determined,
        or which raised a URLError on download.
        """


Spider = Crawler

# class Polly(Crawler):
#    def visit(self, link, source=None):
#        print("visited:", link.url, "from:", link.referrer)
#    def fail(self, link):
#        print("failed:", link.url)
#
# p = Polly(links=["http://nodebox.net/"], domains=["nodebox.net"], delay=5)
# while not p.done:
#    p.crawl(method=DEPTH, cached=True, throttle=5)

# --- CRAWL FUNCTION --------------------------------------------------------------------------------
# Functional approach to crawling.


def crawl(
    links=[],
    domains=[],
    delay=20.0,
    parse=HTMLLinkParser().parse,
    sort=FIFO,
    method=DEPTH,
    **kwargs
):
    """Returns a generator that yields (Link, source)-tuples of visited pages.
    When the crawler is idle, it yields (None, None).
    """
    # The scenarios below defines "idle":
    # - crawl(delay=10, throttle=0)
    #   The crawler will wait 10 seconds before visiting the same subdomain.
    #   The crawler will not throttle downloads, so the next link is visited instantly.
    #   So sometimes (None, None) is returned while it waits for an available subdomain.
    # - crawl(delay=0, throttle=10)
    #   The crawler will wait 10 seconds after each and any visit.
    #   The crawler will not delay before visiting the same subdomain.
    #   So usually a result is returned each crawl.next(), but each call takes 10 seconds.
    # - asynchronous(crawl().next)
    #   AsynchronousRequest.value is set to (Link, source) once AsynchronousRequest.done=True.
    #   The program will not halt in the meantime (i.e., the next crawl is threaded).
    crawler = Crawler(links, domains, delay, parse, sort)
    bind(
        crawler,
        "visit",
        lambda crawler, link, source=None: setattr(crawler, "crawled", (link, source)),
    )  # Define Crawler.visit() on-the-fly.
    while not crawler.done:
        crawler.crawled = (None, None)
        crawler.crawl(method, **kwargs)
        yield crawler.crawled


# for link, source in crawl("http://www.clips.ua.ac.be/", delay=0, throttle=1, cached=False):
#    print(link)

# g = crawl("http://www.clips.ua.ac.be/"")
# for i in range(10):
#    p = asynchronous(g.next)
#    while not p.done:
#        print("zzz...")
#        time.sleep(0.1)
#    link, source = p.value
#    print(link)


#### DOCUMENT PARSER ###############################################################################
# Not to be confused with Document, which is the top-level element in the HTML DOM.


class DocumentParserError(Exception):
    pass


class DocumentParser(object):
    def __init__(self, path, *args, **kwargs):
        """Parses a text document (e.g., .pdf or .docx),
        given as a file path or a string.
        """
        self.content = self._parse(path, *args, **kwargs)

    def _open(self, path):
        """Returns a file-like object with a read() method,
        from the given file path or string.
        """
        if isinstance(path, str) and os.path.exists(path):
            return open(path, "rb")
        if hasattr(path, "read"):
            return path
        return StringIO(path)

    def _parse(self, path, *args, **kwargs):
        """Returns a plaintext Unicode string parsed from the given document."""
        return plaintext(decode_utf8(self.open(path).read()))

    @property
    def string(self):
        return self.content

    def __str__(self):
        return self.content


# --- PDF PARSER ------------------------------------------------------------------------------------
#  Yusuke Shinyama, PDFMiner, http://www.unixuser.org/~euske/python/pdfminer/


class PDFError(DocumentParserError):
    pass


class PDF(DocumentParser):
    def __init__(self, path, output="txt"):
        self.content = self._parse(path, format=output)

    def _parse(self, path, *args, **kwargs):
        # The output is useful for mining but not for display.
        # Alternatively, PDF(format="html") preserves some layout.
        from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
        from pdfminer.pdfpage import PDFPage
        from pdfminer.converter import TextConverter, HTMLConverter
        from pdfminer.layout import LAParams

        try:
            m = PDFResourceManager()
            s = StringIO()
            p = (
                kwargs.get("format", "txt").endswith("html")
                and HTMLConverter
                or TextConverter
            )
            p = p(m, s, codec="utf-8", laparams=LAParams())
            interpreter = PDFPageInterpreter(m, p)
            f = self._open(path)
            for page in PDFPage.get_pages(f, maxpages=0, password=""):
                interpreter.process_page(page)
            f.close()
        except Exception as e:
            raise PDFError(str(e))
        s = s.getvalue()
        s = decode_utf8(s)
        s = s.strip()
        s = re.sub(r"([a-z])\-\n", "\\1", s)  # Hyphenation.
        s = s.replace("\n\n", "<!-- #p -->")  # Paragraphs.
        s = s.replace("\n", " ")
        s = s.replace("<!-- #p -->", "\n\n")
        s = collapse_spaces(s)
        return s


# --- OOXML PARSER ----------------------------------------------------------------------------------
# python-docx, https://github.com/python-openxml/python-docx


class DOCXError(DocumentParserError):
    pass


class DOCX(DocumentParser):
    def _parse(self, path, *args, **kwargs):
        from docx import Document

        document = Document(path)
        try:
            s = [paragraph.text for paragraph in document.paragraphs]
        except Exception as e:
            raise DOCXError(str(e))
        s = "\n\n".join(p for p in s)
        s = decode_utf8(s)
        s = collapse_spaces(s)
        return s


# ---------------------------------------------------------------------------------------------------


def parsepdf(path, *args, **kwargs):
    """Returns the content as a Unicode string from the given .pdf file."""
    return PDF(path, *args, **kwargs).content


def parsedocx(path, *args, **kwargs):
    """Returns the content as a Unicode string from the given .docx file."""
    return DOCX(path, *args, **kwargs).content


def parsehtml(path, *args, **kwargs):
    """Returns the content as a Unicode string from the given .html file."""
    return plaintext(DOM(path, *args, **kwargs).body)


def parsedoc(path, format=None):
    """Returns the content as a Unicode string from the given document (.html., .pdf, .docx)."""
    if isinstance(path, str):
        if format == "pdf" or path.endswith(".pdf"):
            return parsepdf(path)
        if format == "docx" or path.endswith(".docx"):
            return parsedocx(path)
        if format == "html" or path.endswith((".htm", ".html", ".xhtml")):
            return parsehtml(path)
    # Brute-force approach if the format is unknown.
    for f in (parsepdf, parsedocx, parsehtml):
        try:
            return f(path)
        except:
            pass


if __name__ == "__main__":
    from higgins.automation.email import email_utils

    email_id = "0f84cbaf27e24d6d3eb96b734e8f8c776a8fc7250ed479ba252d47a9fc56fcc2"
    # email_id = "e9e027410ae591c3ea94c3d2c67fe7ac05025dd92579e224b1e9abe8dd6ff72d"
    # email_id = "365b47046e7027df274d0a26de33461f99c4cb054e47a6aaa4a6bc791e263585"

    email = email_utils.load_email(email_id)
    # print(email["html"])
    print(email["plain"])

    print("--" * 50)

    # this does better than v2/v3, but not v4 on tables. Maybe we use different
    # parser for tables
    text = plaintext(email["html"])
    print(text)

    # print("--" * 50)
    # print(email_utils.parse_html_v4(email["html"]))

    # print("--" * 50)
    # print(email_utils.parse_html_v3(email["html"]))
